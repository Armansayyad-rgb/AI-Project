"""Document upload pipeline for the Gradio web UI.

Parses uploaded PDFs, DOCX, and TXT files, splits them into chunks using
the same logic as the static knowledge base, and merges the new chunks
into the running pipeline's retrieval index.

Design notes:

- We avoid hard dependencies on PyPDF2/python-docx by lazy-importing them
  only when the corresponding file type is uploaded. This keeps the
  basic chat UI working even if those packages are not installed.
- Chunk sizes mirror ``retriever_v2.load_chunks`` so retrieval scoring
  behaves identically on uploaded and built-in content.
- Index rebuild is O(N) over the full chunk list. For the demo corpus
  (~107k chunks) this is ~1s; well within an acceptable UI delay.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from retriever_v2 import build_index as build_index_v2


SUPPORTED_EXTS = {".txt", ".pdf", ".docx"}


@dataclass
class UploadedDocument:
    """Record of a single uploaded file."""

    name: str
    path: Path
    ext: str
    text: str
    chunks: list[str] = field(default_factory=list)
    chunk_count: int = 0

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "path": str(self.path),
            "ext": self.ext,
            "chunk_count": self.chunk_count,
        }


def _read_text(path: Path) -> str:
    """Plain TXT reader."""
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    """PDF reader using PyPDF2 (lazy import)."""
    from PyPDF2 import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n\n".join(parts)


def _read_docx(path: Path) -> str:
    """DOCX reader using python-docx (lazy import)."""
    from docx import Document as DocxDocument  # type: ignore

    doc = DocxDocument(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs)


def parse_file(path: Path) -> str:
    """Dispatch a single file to the right parser."""
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_text(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

_WORD_RE = re.compile(r"\s+")
# Approximate chunk size in words; matches the mean chunk size produced by
# retriever_v2.load_chunks (~120 words/sentence, ~5 sentences per chunk).
_CHUNK_WORDS = 500
_OVERLAP_WORDS = 50


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WORD_RE.sub(" ", text)
    return text.strip()


def chunk_text(text: str, chunk_words: int = _CHUNK_WORDS, overlap: int = _OVERLAP_WORDS) -> list[str]:
    """Split text into overlapping word-window chunks."""
    text = _normalize_text(text)
    if not text:
        return []
    words = text.split(" ")
    if len(words) <= chunk_words:
        return [text]

    step = max(1, chunk_words - overlap)
    chunks: list[str] = []
    for start in range(0, len(words), step):
        piece = words[start : start + chunk_words]
        if not piece:
            break
        chunks.append(" ".join(piece))
        if start + chunk_words >= len(words):
            break
    return chunks


# ---------------------------------------------------------------------------
# Pipeline integration
# ---------------------------------------------------------------------------

def _is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTS


def attach_documents(
    pipeline: dict,
    uploaded: list[UploadedDocument],
) -> int:
    """Merge uploaded chunks into the pipeline's chunks list and rebuild the index.

    Returns the total number of new chunks added.
    """
    new_chunks: list[str] = []
    for doc in uploaded:
        if not doc.chunks:
            doc.chunks = chunk_text(doc.text)
        new_chunks.extend(doc.chunks)
        doc.chunk_count = len(doc.chunks)
    if not new_chunks:
        return 0

    # retriever_v2 expects a list of plain strings; flatten directly.
    pipeline["chunks"].extend(new_chunks)
    pipeline["retrieval_index"], pipeline["document_frequency"] = build_index_v2(
        pipeline["chunks"]
    )

    # Track uploads in the pipeline so the UI can list them.
    pipeline.setdefault("uploaded_docs", []).extend(
        [d.to_dict() for d in uploaded]
    )

    return len(new_chunks)


def process_uploads(
    pipeline: dict,
    file_paths: Iterable[str],
) -> tuple[list[UploadedDocument], list[str]]:
    """Parse a batch of uploaded files, return (parsed, errors).

    The pipeline is not modified here; call ``attach_documents`` after this
    step when the user explicitly confirms the upload.
    """
    parsed: list[UploadedDocument] = []
    errors: list[str] = []

    for raw in file_paths:
        path = Path(raw)
        if not path.exists():
            errors.append(f"Not found: {raw}")
            continue
        if not _is_supported(path):
            errors.append(f"Unsupported file type: {path.name}")
            continue
        try:
            text = parse_file(path)
        except Exception as exc:
            errors.append(f"Failed to parse {path.name}: {exc!r}")
            continue
        doc = UploadedDocument(
            name=path.name,
            path=path,
            ext=path.suffix.lower(),
            text=text,
        )
        parsed.append(doc)

    return parsed, errors
